"""
CV Generator Module
Handles CV generation with AI enhancements, ATS optimization, and PDF/Word export.
"""
from typing import Dict, Any, Optional, List
from sqlalchemy.orm import Session
import sys
from pathlib import Path

# Add backend directory to path for imports
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from app.services.ai_service import AIService
from app.db.models import CV, User, Credential
from app.utils.logger import logger


class CVGenerator:
    """CV generation service with AI enhancements."""
    
    def __init__(self):
        self.ai_service = AIService()
    
    def generate_cv_with_market_analysis(
        self,
        user_id: int,
        personal_info: Dict[str, Any],
        experience: List[Dict[str, Any]],
        education: List[Dict[str, Any]],
        skills: Dict[str, Any],
        industry: Optional[str] = None,
        awards: List[Dict[str, Any]] = None,
        publications: List[Dict[str, Any]] = None,
        projects: List[Dict[str, Any]] = None,
        memberships: List[Dict[str, Any]] = None,
        job_id: Optional[int] = None,
        photo_url: Optional[str] = None,
        db: Session = None
    ) -> Dict[str, Any]:
        """
        Generate CV with market analysis FIRST, then AI enhancements.
        This ensures suggestions are based on real market trends.
        """
        logger.info(f"Generating CV with market analysis for user {user_id}")
        
        # STEP 1: Analyze job market for the user's industry
        market_analysis = {}
        if db and industry:
            market_analysis = self.ai_service.analyze_job_market(db, sector=industry)
            logger.info(f"Market analysis: {len(market_analysis.get('trending_skills', []))} trending skills identified")
        
        # STEP 2: Generate CV (now informed by market analysis)
        cv_result = self.generate_cv(
            user_id=user_id,
            personal_info=personal_info,
            experience=experience,
            education=education,
            skills=skills,
            awards=awards,
            publications=publications,
            projects=projects,
            memberships=memberships,
            job_id=job_id,
            photo_url=photo_url,
            db=db
        )
        
        # STEP 3: Add market analysis to CV result
        cv_result["market_analysis"] = market_analysis
        
        return cv_result
    
    def generate_cv(
        self,
        user_id: int,
        personal_info: Dict[str, Any],
        experience: List[Dict[str, Any]],
        education: List[Dict[str, Any]],
        skills: Dict[str, Any],
        awards: List[Dict[str, Any]] = None,
        publications: List[Dict[str, Any]] = None,
        projects: List[Dict[str, Any]] = None,
        memberships: List[Dict[str, Any]] = None,
        job_id: Optional[int] = None,
        photo_url: Optional[str] = None,
        db: Session = None,
        original_file_url: Optional[str] = None,
        original_file_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a professional CV with AI enhancements.
        
        Args:
            user_id: User ID
            personal_info: Personal information dict
            experience: List of work experience entries
            education: List of education entries
            skills: Skills dictionary (Europass format)
            awards: Optional awards list
            publications: Optional publications list
            projects: Optional projects list
            memberships: Optional memberships list
            job_id: Optional job ID to tailor CV
            photo_url: Optional photo URL
            db: Database session
            
        Returns:
            Generated CV data with AI enhancements
        """
        logger.info(f"Generating CV for user {user_id}")
        
        # Get user
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise ValueError(f"User {user_id} not found")
        
        # Prepare user data
        user_data = {
            "full_name": personal_info.get("first_name", "") + " " + personal_info.get("surname", "") or user.full_name,
            "email": personal_info.get("email") or user.email,
            "wallet_address": user.wallet_address,
            "role": user.role.value,
            "photo_url": photo_url,
            **personal_info
        }
        
        # Generate CV with AI (certificates removed - using education list)
        cv_json = self.ai_service.generate_cv(
            user_data=user_data,
            certificates=[],  # Certificates removed - use education list instead
            experience=experience,
            education=education,
            skills=skills,
            awards=awards or [],
            publications=publications or [],
            projects=projects or [],
            memberships=memberships or [],
            job_id=job_id,
            db=db
        )
        
        # CRITICAL: Create credentials (single source of truth) before saving CV
        from app.services.trust_service import CredentialService
        from app.db.models import CredentialType, CredentialSource
        from datetime import datetime
        
        credential_service = CredentialService()
        credentials_created = []
        
        # Create education credentials
        for edu in education:
            if not isinstance(edu, dict):
                continue
            degree = edu.get("degree") or edu.get("qualification") or ""
            institution = edu.get("institution") or edu.get("school") or ""
            
            if not degree:
                continue
            
            try:
                # Check if credential already exists
                existing = (
                    db.query(Credential)
                    .filter(
                        Credential.user_id == user_id,
                        Credential.type == CredentialType.EDUCATION,
                        Credential.title == degree,
                        Credential.organization == institution,
                    )
                    .first()
                )
                
                if not existing:
                    start_date = None
                    if edu.get("start_date"):
                        try:
                            start_date = datetime.fromisoformat(edu["start_date"].replace("Z", "+00:00"))
                        except:
                            pass
                    
                    end_date = None
                    if edu.get("graduation_year"):
                        try:
                            end_date = datetime(int(edu["graduation_year"]), 1, 1)
                        except:
                            pass
                    elif edu.get("end_date"):
                        try:
                            end_date = datetime.fromisoformat(edu["end_date"].replace("Z", "+00:00"))
                        except:
                            pass
                    
                    credential = credential_service.create_credential(
                        db=db,
                        user_id=user_id,
                        credential_type=CredentialType.EDUCATION,
                        title=degree,
                        organization=institution,
                        start_date=start_date,
                        end_date=end_date,
                        description=edu.get("field_of_study", ""),
                        source=CredentialSource.USER_INPUT,
                    )
                    credentials_created.append(credential.id)
            except Exception as e:
                logger.error(f"Error creating education credential: {e}")
                continue
        
        # Create employment credentials
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            job_title = exp.get("job_title") or exp.get("position") or ""
            company = exp.get("company") or exp.get("employer") or ""
            
            if not job_title and not company:
                continue
            
            try:
                # Check if credential already exists
                existing = (
                    db.query(Credential)
                    .filter(
                        Credential.user_id == user_id,
                        Credential.type == CredentialType.EMPLOYMENT,
                        Credential.title == job_title,
                        Credential.organization == company,
                    )
                    .first()
                )
                
                if not existing:
                    start_date = None
                    if exp.get("start_date"):
                        try:
                            start_date = datetime.fromisoformat(exp["start_date"].replace("Z", "+00:00"))
                        except:
                            pass
                    
                    end_date = None
                    if exp.get("end_date") and exp.get("end_date") != "Present":
                        try:
                            end_date = datetime.fromisoformat(exp["end_date"].replace("Z", "+00:00"))
                        except:
                            pass
                    
                    credential = credential_service.create_credential(
                        db=db,
                        user_id=user_id,
                        credential_type=CredentialType.EMPLOYMENT,
                        title=job_title or "Employee",
                        organization=company,
                        start_date=start_date,
                        end_date=end_date,
                        description=exp.get("description", ""),
                        source=CredentialSource.USER_INPUT,
                    )
                    credentials_created.append(credential.id)
            except Exception as e:
                logger.error(f"Error creating employment credential: {e}")
                continue
        
        logger.info(f"Created {len(credentials_created)} credentials for user {user_id}")
        
        # CRITICAL: Check if CV already exists - if so, preserve stored_job_matches, cv_domain, parsed_skills, job_search
        # These are set during Quick Upload and MUST be preserved when CV is saved/updated via other endpoints
        existing_cv = db.query(CV).filter(CV.user_id == user_id).order_by(CV.created_at.desc()).first()
        
        # CRITICAL: Preserve ALL job matching data if CV already exists
        # This ensures jobs from Quick Upload are NEVER lost when CV Builder updates the CV
        preserved_data = {}
        if existing_cv and existing_cv.json_content and isinstance(existing_cv.json_content, dict):
            # Preserve ALL job-related keys
            job_related_keys = [
                "stored_job_matches", 
                "cv_domain", 
                "parsed_skills", 
                "job_search", 
                "job_match_keywords", 
                "pdf_url",
                "job_matches",  # Also preserve if exists
                "last_job_search",  # Also preserve if exists
            ]
            for key in job_related_keys:
                if key in existing_cv.json_content:
                    preserved_data[key] = existing_cv.json_content[key]
                    logger.info(f"[CV GENERATOR] Preserving {key} from existing CV {existing_cv.id} (type: {type(existing_cv.json_content[key])}, length: {len(existing_cv.json_content[key]) if isinstance(existing_cv.json_content[key], (list, dict)) else 'N/A'})")
            
            # CRITICAL: Ensure stored_job_matches is always a list (never None)
            if "stored_job_matches" in preserved_data:
                if preserved_data["stored_job_matches"] is None:
                    logger.warning(f"[CV GENERATOR] stored_job_matches was None - converting to empty list")
                    preserved_data["stored_job_matches"] = []
                elif not isinstance(preserved_data["stored_job_matches"], list):
                    logger.warning(f"[CV GENERATOR] stored_job_matches was {type(preserved_data['stored_job_matches'])} - converting to list")
                    preserved_data["stored_job_matches"] = []
        
        # Merge preserved data into new CV JSON
        final_cv_json = {**cv_json, **preserved_data}
        
        # Save CV (still needed for skills, summary, ATS score, etc.)
        if existing_cv:
            # Update existing CV instead of creating new one
            existing_cv.json_content = final_cv_json
            existing_cv.ai_score = cv_json.get("ai_score")
            if photo_url:
                existing_cv.photo_url = photo_url
            if original_file_url:
                existing_cv.original_file_url = original_file_url
            if original_file_name:
                existing_cv.original_file_name = original_file_name
            
            # CRITICAL: Flag json_content as modified so SQLAlchemy detects the change
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(existing_cv, "json_content")
            
            db.commit()
            db.refresh(existing_cv)
            cv = existing_cv
            logger.info(f"[CV GENERATOR] Updated existing CV {cv.id}, preserved {len(preserved_data)} job matching fields")
        else:
            # Create new CV
            cv = CV(
                user_id=user_id,
                json_content=final_cv_json,
                ai_score=cv_json.get("ai_score"),
                photo_url=photo_url,
                original_file_url=original_file_url,
                original_file_name=original_file_name
            )
            db.add(cv)
            db.commit()
            db.refresh(cv)
            logger.info(f"[CV GENERATOR] Created new CV {cv.id}")
        
        logger.info(f"CV generated with score: {cv.ai_score}, {len(credentials_created)} credentials created")
        return {
            "id": cv.id,
            "user_id": cv.user_id,
            "json_content": cv.json_content,
            "ai_score": cv.ai_score,
            "photo_url": cv.photo_url,
            "credentials_created": len(credentials_created),
            "created_at": cv.created_at.isoformat() if cv.created_at else None,
            "updated_at": cv.updated_at.isoformat() if cv.updated_at else None,
        }
    
    def get_cv(self, user_id: int, db: Session) -> Optional[Dict[str, Any]]:
        """
        Get CV data for a user.
        
        Now reads from credentials table (single source of truth) while maintaining
        backward compatibility with existing CV JSON data.
        """
        from app.services.trust_service import CredentialService
        from app.db.models import CredentialType, VerificationStatus
        
        credential_service = CredentialService()
        
        # Get CV record (for template, PDF URL, etc.)
        cv = db.query(CV).filter(CV.user_id == user_id).order_by(CV.created_at.desc()).first()
        
        # Get credentials (single source of truth)
        education_creds = credential_service.get_user_credentials(
            db, user_id, CredentialType.EDUCATION
        )
        
        employment_creds = credential_service.get_user_credentials(
            db, user_id, CredentialType.EMPLOYMENT
        )
        
        startup_role_creds = credential_service.get_user_credentials(
            db, user_id, CredentialType.STARTUP_ROLE
        )
        
        # Convert credentials to CV format
        education = [
            {
                "degree": cred.title,
                "institution": cred.organization or "",
                "start_date": cred.start_date.isoformat() if cred.start_date else None,
                "graduation_year": cred.end_date.year if cred.end_date else None,
                "end_date": cred.end_date.isoformat() if cred.end_date else None,
                "field_of_study": cred.description or "",
                "verified": cred.verification_status == VerificationStatus.VERIFIED,
                "credential_id": cred.id,
            }
            for cred in education_creds
        ]
        
        experience = [
            {
                "job_title": cred.title,
                "company": cred.organization or "",
                "start_date": cred.start_date.isoformat() if cred.start_date else None,
                "end_date": cred.end_date.isoformat() if cred.end_date else "Present",
                "description": cred.description or "",
                "verified": cred.verification_status == VerificationStatus.VERIFIED,
                "credential_id": cred.id,
            }
            for cred in employment_creds + startup_role_creds  # Include startup roles as experience
        ]
        
        # Build CV data structure from credentials
        cv_data = {
            "education": education,
            "experience": experience,
            "work_experience": experience,  # Backward compatibility
        }
        
        # Check if there's any actual data (CV record OR credentials)
        has_cv_record = cv is not None
        has_credentials = len(education) > 0 or len(experience) > 0
        
        # If there's no CV record and no credentials, return None (nothing to show)
        if not has_cv_record and not has_credentials:
            return None
        
        # Merge with existing CV JSON if it exists (for skills, summary, personal_info, etc.)
        if cv and cv.json_content:
            existing_data = cv.json_content.copy()
            
            # CRITICAL: Extract personal_info from various possible locations
            # It might be at top level, nested, or in optimized_cv/original_cv
            personal_info = existing_data.get("personal_info", {})
            if not personal_info or not isinstance(personal_info, dict):
                # Try to extract from nested structures
                if existing_data.get("optimized_cv") and isinstance(existing_data["optimized_cv"], dict):
                    personal_info = existing_data["optimized_cv"].get("personal_info", {})
                elif existing_data.get("original_cv") and isinstance(existing_data["original_cv"], dict):
                    personal_info = existing_data["original_cv"].get("personal_info", {})
                elif existing_data.get("original_cv_data") and isinstance(existing_data["original_cv_data"], dict):
                    personal_info = existing_data["original_cv_data"].get("personal_info", {})
            
            # Override education/experience with credential-based data when available
            # If no credentials were created (e.g., creation error), keep stored CV sections intact
            # BUT also extract from nested optimized_cv/original_cv if top-level is empty
            if education:
                existing_data["education"] = education
            elif not existing_data.get("education"):
                # Top-level education is missing or empty - check nested structures
                existing_data["education"] = (
                    existing_data.get("optimized_cv", {}).get("education") or
                    existing_data.get("original_cv", {}).get("education") or
                    existing_data.get("original_cv_data", {}).get("education") or
                    []
                )

            if experience:
                existing_data["experience"] = experience
                existing_data["work_experience"] = experience
            elif not existing_data.get("experience") and not existing_data.get("work_experience"):
                # Top-level experience is missing or empty - check nested structures
                nested_exp = (
                    existing_data.get("optimized_cv", {}).get("experience") or
                    existing_data.get("original_cv", {}).get("experience") or
                    existing_data.get("original_cv_data", {}).get("experience") or
                    existing_data.get("optimized_cv", {}).get("work_experience") or
                    existing_data.get("original_cv", {}).get("work_experience") or
                    []
                )
                existing_data["experience"] = nested_exp
                existing_data["work_experience"] = nested_exp
            else:
                # Ensure keys exist without clobbering existing CV content
                if "experience" not in existing_data:
                    existing_data["experience"] = []
                if "work_experience" not in existing_data:
                    existing_data["work_experience"] = existing_data.get("experience", [])
            
            # Ensure personal_info is preserved
            if personal_info and isinstance(personal_info, dict):
                existing_data["personal_info"] = personal_info
            
            cv_data = existing_data
            
            # Ensure ATS score is included
            if cv.ai_score is not None and "ats_score" not in cv_data:
                cv_data["ats_score"] = cv.ai_score
            if "ats_grade" not in cv_data and cv.ai_score is not None:
                if cv.ai_score >= 90:
                    cv_data["ats_grade"] = "A+"
                elif cv.ai_score >= 80:
                    cv_data["ats_grade"] = "A"
                elif cv.ai_score >= 70:
                    cv_data["ats_grade"] = "B"
                elif cv.ai_score >= 60:
                    cv_data["ats_grade"] = "C"
                else:
                    cv_data["ats_grade"] = "D"
        elif not cv:
            # No CV record exists, return credential-based data only
            cv_data = {
                "education": education,
                "experience": experience,
                "work_experience": experience,
            }
        
        # Extract PDF URL if it exists
        pdf_url = None
        if cv and cv.json_content:
            pdf_url = cv.json_content.get("pdf_url")
        
        # Normalize data structure for frontend compatibility
        # Ensure experience and work_experience both exist
        if "work_experience" in cv_data and "experience" not in cv_data:
            cv_data["experience"] = cv_data["work_experience"]
        elif "experience" in cv_data and "work_experience" not in cv_data:
            cv_data["work_experience"] = cv_data["experience"]
        
        # Ensure skills and personal_skills both exist
        if "personal_skills" in cv_data and "skills" not in cv_data:
            cv_data["skills"] = cv_data["personal_skills"]
        elif "skills" in cv_data and "personal_skills" not in cv_data:
            cv_data["personal_skills"] = cv_data["skills"]
        
        # CRITICAL: Ensure personal_info exists with default structure for CV Editor
        # If personal_info is missing or empty, populate from user data
        if "personal_info" not in cv_data or not isinstance(cv_data["personal_info"], dict):
            cv_data["personal_info"] = {}
        
        # Get user data to populate missing fields
        user = db.query(User).filter(User.id == user_id).first() if db else None
        
        # Ensure personal_info has required fields - populate from user if missing
        if not cv_data["personal_info"].get("full_name"):
            # Try to get from existing data first
            full_name = (
                cv_data["personal_info"].get("full_name") or
                f"{cv_data['personal_info'].get('first_name', '')} {cv_data['personal_info'].get('surname', '')}".strip() or
                (user.full_name if user else "")
            )
            cv_data["personal_info"]["full_name"] = full_name
        
        if not cv_data["personal_info"].get("email"):
            cv_data["personal_info"]["email"] = (
                cv_data["personal_info"].get("email") or
                (user.email if user else "")
            )
        
        if not cv_data["personal_info"].get("phone"):
            cv_data["personal_info"]["phone"] = cv_data["personal_info"].get("phone") or ""
        
        if not cv_data["personal_info"].get("location"):
            cv_data["personal_info"]["location"] = (
                cv_data["personal_info"].get("location") or
                cv_data["personal_info"].get("address") or
                ""
            )
        
        # Also ensure first_name and surname exist for backward compatibility
        if not cv_data["personal_info"].get("first_name") and cv_data["personal_info"].get("full_name"):
            name_parts = cv_data["personal_info"]["full_name"].split()
            cv_data["personal_info"]["first_name"] = name_parts[0] if name_parts else ""
            cv_data["personal_info"]["surname"] = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
        
        # Log what we found for debugging
        logger.info(f"[GET_CV] personal_info extracted: full_name={cv_data['personal_info'].get('full_name', 'MISSING')}, email={cv_data['personal_info'].get('email', 'MISSING')}, has_data={bool(cv_data['personal_info'].get('full_name') or cv_data['personal_info'].get('email'))}")
        
        # CRITICAL: Ensure summary exists (at top level and in json_content for CV Editor)
        # Extract summary from various possible locations
        summary = cv_data.get("summary", "")
        if not summary:
            # Try to get from nested structures
            if isinstance(cv_data.get("optimized_cv"), dict):
                summary = cv_data["optimized_cv"].get("summary", "")
            elif isinstance(cv_data.get("original_cv"), dict):
                summary = cv_data["original_cv"].get("summary", "")
            elif isinstance(cv_data.get("original_cv_data"), dict):
                summary = cv_data["original_cv_data"].get("summary", "")
        
        cv_data["summary"] = summary
        
        # Log what we found for debugging
        logger.info(f"[GET_CV] Summary extracted: length={len(summary)}, has_content={bool(summary)}")
        
        # CRITICAL: Ensure personal_skills exists with proper structure
        if "personal_skills" not in cv_data:
            cv_data["personal_skills"] = {}
        if not isinstance(cv_data["personal_skills"], dict):
            cv_data["personal_skills"] = {}
        if "job_related_skills" not in cv_data["personal_skills"]:
            cv_data["personal_skills"]["job_related_skills"] = []
        if not isinstance(cv_data["personal_skills"]["job_related_skills"], list):
            cv_data["personal_skills"]["job_related_skills"] = []
        
        # Ensure education and experience are lists (not None)
        if "education" not in cv_data:
            cv_data["education"] = []
        if not isinstance(cv_data["education"], list):
            cv_data["education"] = []
        if "experience" not in cv_data:
            cv_data["experience"] = []
        if not isinstance(cv_data["experience"], list):
            cv_data["experience"] = []
        
        # Final validation: Log what we're returning
        logger.info(f"[GET_CV] Returning CV data for user {user_id}: "
                   f"has_personal_info={bool(cv_data.get('personal_info'))}, "
                   f"personal_info_keys={list(cv_data.get('personal_info', {}).keys()) if cv_data.get('personal_info') else []}, "
                   f"personal_info_full_name={cv_data.get('personal_info', {}).get('full_name', 'MISSING')}, "
                   f"personal_info_email={cv_data.get('personal_info', {}).get('email', 'MISSING')}, "
                   f"has_summary={bool(cv_data.get('summary'))}, "
                   f"summary_length={len(cv_data.get('summary', ''))}, "
                   f"has_skills={bool(cv_data.get('personal_skills'))}, "
                   f"json_content_keys={list(cv_data.keys())[:15]}")
        
        return {
            "id": cv.id if cv else None,
            "user_id": user_id,
            "json_content": cv_data,
            "summary": cv_data.get("summary", ""),  # Also include at top level for CV Editor compatibility
            "ai_score": cv.ai_score if cv else None,
            "photo_url": cv.photo_url if cv else None,
            "pdf_url": pdf_url,  # Include PDF URL at top level for easy access
            "created_at": cv.created_at.isoformat() if cv and cv.created_at else None,
            "updated_at": cv.updated_at.isoformat() if cv and cv.updated_at else None,
        }
    
    def export_to_pdf(self, cv_data: Dict[str, Any]) -> bytes:
        """
        Export CV to PDF format using reportlab.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib import colors
            from io import BytesIO
            
            logger.info("Exporting CV to PDF")
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=12,
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=6,
                spaceBefore=12,
            )
            
            # Personal Info (fallback through common shapes)
            personal_info = cv_data.get("personal_info", {}) or cv_data.get("json_content", {}).get("personal_info", {})
            if personal_info:
                name = personal_info.get("full_name", "CV")
                email = personal_info.get("email", "")
                phone = personal_info.get("phone", "")
                address = personal_info.get("address", "")
                
                story.append(Paragraph(name, title_style))
                contact_info = []
                if email:
                    contact_info.append(email)
                if phone:
                    contact_info.append(phone)
                if address:
                    contact_info.append(address)
                if contact_info:
                    story.append(Paragraph(" | ".join(contact_info), styles['Normal']))
                story.append(Spacer(1, 0.2*inch))

            # Summary
            summary = (
                cv_data.get("summary")
                or cv_data.get("json_content", {}).get("summary")
                or personal_info.get("summary")
            )
            if summary:
                story.append(Paragraph("Professional Summary", heading_style))
                story.append(Paragraph(summary, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Experience (check for non-empty lists explicitly)
            experience = cv_data.get("experience") or cv_data.get("work_experience")
            if not experience or (isinstance(experience, list) and len(experience) == 0):
                json_content = cv_data.get("json_content", {})
                experience = json_content.get("experience") or json_content.get("work_experience")
                # If still empty, check nested optimized_cv/original_cv
                if not experience or (isinstance(experience, list) and len(experience) == 0):
                    experience = (
                        json_content.get("optimized_cv", {}).get("experience") or
                        json_content.get("optimized_cv", {}).get("work_experience") or
                        json_content.get("original_cv", {}).get("experience") or
                        json_content.get("original_cv", {}).get("work_experience") or
                        []
                    )
            if experience and len(experience) > 0:
                story.append(Paragraph("Experience", heading_style))
                for exp in experience:
                    title = exp.get("job_title", "")
                    company = exp.get("company", "")
                    period = exp.get("period", "")
                    if not period:
                        start = exp.get("start_date")
                        end = exp.get("end_date") or "Present"
                        if start:
                            period = f"{start} - {end}"
                    description = exp.get("description", "")
                    
                    exp_text = f"<b>{title}</b>"
                    if company:
                        exp_text += f" - {company}"
                    if period:
                        exp_text += f" ({period})"
                    story.append(Paragraph(exp_text, styles['Normal']))
                    if description:
                        story.append(Paragraph(description, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Education (check for non-empty lists explicitly)
            education = cv_data.get("education")
            if not education or (isinstance(education, list) and len(education) == 0):
                json_content = cv_data.get("json_content", {})
                education = json_content.get("education")
                # If still empty, check nested optimized_cv/original_cv
                if not education or (isinstance(education, list) and len(education) == 0):
                    education = (
                        json_content.get("optimized_cv", {}).get("education") or
                        json_content.get("original_cv", {}).get("education") or
                        []
                    )
            if education and len(education) > 0:
                story.append(Paragraph("Education", heading_style))
                for edu in education:
                    degree = edu.get("degree", "")
                    institution = edu.get("institution", "")
                    period = edu.get("period", "")
                    if not period:
                        start = edu.get("start_date")
                        end = edu.get("end_date") or edu.get("graduation_year")
                        if start or end:
                            period = f"{start or ''} - {end or ''}".strip(" -")
                    
                    edu_text = f"<b>{degree}</b>"
                    if institution:
                        edu_text += f" - {institution}"
                    if period:
                        edu_text += f" ({period})"
                    story.append(Paragraph(edu_text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Skills
            skills = (
                cv_data.get("skills")
                or cv_data.get("personal_skills")
                or cv_data.get("json_content", {}).get("skills")
                or cv_data.get("json_content", {}).get("personal_skills")
                or {}
            )
            if skills:
                story.append(Paragraph("Skills", heading_style))
                skills_text = []
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if isinstance(skill_list, list):
                            skills_text.append(f"{category}: {', '.join(skill_list)}")
                        else:
                            skills_text.append(f"{category}: {skill_list}")
                else:
                    skills_text.append(str(skills))
                
                if skills_text:
                    story.append(Paragraph(" | ".join(skills_text), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("reportlab not installed. Install with: pip install reportlab")
            raise NotImplementedError("PDF export requires reportlab. Install with: pip install reportlab")
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}", exc_info=True)
            raise
    
    def export_to_docx(self, cv_data: Dict[str, Any]) -> bytes:
        """Alias for export_to_word - exports CV to DOCX format."""
        return self.export_to_word(cv_data)
    
    def export_to_word(self, cv_data: Dict[str, Any]) -> bytes:
        """
        Export CV to Word format using python-docx.
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            logger.info("Exporting CV to Word")
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Personal Info (fallback through common shapes - same as PDF)
            personal_info = cv_data.get("personal_info", {}) or cv_data.get("json_content", {}).get("personal_info", {})
            if personal_info:
                name = personal_info.get("full_name", "CV")
                if not name:
                    name = f"{personal_info.get('first_name', '')} {personal_info.get('surname', '')}".strip() or "CV"
                email = personal_info.get("email", "")
                phone = personal_info.get("phone", "") or personal_info.get("mobile", "")
                address = personal_info.get("address", "") or personal_info.get("location", "")
                
                # Name as title
                title = doc.add_heading(name, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0] if title.runs else None
                if title_run:
                    title_run.font.size = Pt(24)
                    title_run.font.bold = True
                
                # Contact info
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_info = []
                if email:
                    contact_info.append(email)
                if phone:
                    contact_info.append(phone)
                if address:
                    contact_info.append(address)
                if contact_info:
                    contact_para.add_run(" | ".join(contact_info))
                
                doc.add_paragraph()  # Spacing
            
            # Summary (same as PDF)
            summary = (
                cv_data.get("summary")
                or cv_data.get("json_content", {}).get("summary")
                or personal_info.get("summary")
            )
            if summary:
                doc.add_heading('PROFESSIONAL SUMMARY', level=1)
                # Strip markdown if present
                import re
                summary_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', summary)
                summary_clean = re.sub(r'\*(.*?)\*', r'\1', summary_clean)
                doc.add_paragraph(summary_clean)
                doc.add_paragraph()  # Spacing
            
            # Experience (check for non-empty lists - same as PDF)
            experience = cv_data.get("experience") or cv_data.get("work_experience")
            if not experience or (isinstance(experience, list) and len(experience) == 0):
                json_content = cv_data.get("json_content", {})
                experience = json_content.get("experience") or json_content.get("work_experience")
                if not experience or (isinstance(experience, list) and len(experience) == 0):
                    experience = (
                        json_content.get("optimized_cv", {}).get("experience") or
                        json_content.get("optimized_cv", {}).get("work_experience") or
                        json_content.get("original_cv", {}).get("experience") or
                        json_content.get("original_cv", {}).get("work_experience") or
                        []
                    )
            
            if experience and len(experience) > 0:
                doc.add_heading('WORK EXPERIENCE', level=1)
                for exp in experience:
                    title = exp.get("job_title", "") or exp.get("position", "")
                    company = exp.get("company", "") or exp.get("employer", "")
                    period = exp.get("period", "")
                    if not period:
                        start = exp.get("start_date", "")
                        end = exp.get("end_date", "Present")
                        if start:
                            period = f"{start} - {end}"
                    description = exp.get("description", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{title}").bold = True
                    if company:
                        p.add_run(f" - {company}")
                    if period:
                        p.add_run(f" ({period})")
                    
                    if description:
                        # Clean markdown
                        import re
                        desc_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', description)
                        desc_clean = re.sub(r'\*(.*?)\*', r'\1', desc_clean)
                        doc.add_paragraph(desc_clean)
            
            # Education (same as PDF)
            education = cv_data.get("education", [])
            if not education or (isinstance(education, list) and len(education) == 0):
                json_content = cv_data.get("json_content", {})
                education = json_content.get("education", [])
            
            if education and len(education) > 0:
                doc.add_heading('EDUCATION', level=1)
                for edu in education:
                    degree = edu.get("degree", "") or edu.get("qualification", "")
                    institution = edu.get("institution", "") or edu.get("school", "")
                    period = edu.get("period", "")
                    if not period:
                        period = edu.get("graduation_year", "") or edu.get("end_date", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{degree}").bold = True
                    if institution:
                        p.add_run(f" - {institution}")
                    if period:
                        p.add_run(f" ({period})")
            
            # Skills (same as PDF)
            skills = cv_data.get("skills", {}) or cv_data.get("personal_skills", {})
            if not skills:
                json_content = cv_data.get("json_content", {})
                skills = json_content.get("skills", {}) or json_content.get("personal_skills", {})
            
            if skills:
                doc.add_heading('SKILLS & COMPETENCIES', level=1)
                if isinstance(skills, dict):
                    # Handle different skill formats
                    job_skills = skills.get("job_related_skills", []) or skills.get("technical_skills", []) or skills.get("technical", [])
                    computer_skills = skills.get("computer_skills", []) or skills.get("programming_skills", [])
                    soft_skills = skills.get("soft", []) or skills.get("social_skills", [])
                    languages = skills.get("languages", [])
                    
                    if job_skills:
                        p = doc.add_paragraph()
                        p.add_run("Technical Skills: ").bold = True
                        p.add_run(", ".join(job_skills))
                    
                    if computer_skills:
                        p = doc.add_paragraph()
                        p.add_run("Programming & Tools: ").bold = True
                        p.add_run(", ".join(computer_skills))
                    
                    if soft_skills:
                        p = doc.add_paragraph()
                        p.add_run("Soft Skills: ").bold = True
                        p.add_run(", ".join(soft_skills))
                    
                    if languages:
                        p = doc.add_paragraph()
                        p.add_run("Languages: ").bold = True
                        p.add_run(", ".join(languages))
                else:
                    doc.add_paragraph(str(skills))
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise NotImplementedError("Word export requires python-docx. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}", exc_info=True)
            raise

